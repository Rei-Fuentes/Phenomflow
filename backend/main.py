from fastapi import FastAPI, Depends, HTTPException
from sqlalchemy.orm import Session
from database import SessionLocal, engine, Base
from models import Analysis
from service import analyze_text, perform_comparative_analysis, analyze_with_pipeline, synthesize_structure
from qdpx_parser import extract_codes_from_qdpx
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware

Base.metadata.create_all(bind=engine)

app = FastAPI(title="PhenomFlow API", description="API for Phenomenological Analysis")

# CORS setup for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

from typing import Optional, List, Dict, Any

class ResearchContext(BaseModel):
    research_question: Optional[str] = None
    study_objective: Optional[str] = None
    phenomenological_approach: Optional[str] = None
    participant_context: Optional[str] = None
    interview_type: Optional[str] = None
    interview_timing: Optional[str] = None

class AnalysisRequest(BaseModel):
    text: str
    context: Optional[ResearchContext] = None
    custom_codes: Optional[List[Dict[str, Any]]] = None

class AnalysisResponse(BaseModel):
    id: int
    input_text: str
    result: str

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

@app.get("/")
def read_root():
    return {"message": "Welcome to PhenomFlow API. Use /analyze to perform analysis."}

@app.post("/analyze", response_model=AnalysisResponse)
def create_analysis(request: AnalysisRequest, db: Session = Depends(get_db)):
    # Perform analysis
    try:
        result_text = analyze_text(request.text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
    # Save to DB
    db_analysis = Analysis(input_text=request.text, result=result_text)
    db.add(db_analysis)
    db.commit()
    db.refresh(db_analysis)
    
    return AnalysisResponse(id=db_analysis.id, input_text=db_analysis.input_text, result=db_analysis.result)

@app.post("/analyze/enhanced")
def create_enhanced_analysis(request: AnalysisRequest):
    """
    Enhanced analysis with 5-phase pipeline returning structured JSON data for visualizations
    """
    try:
        result_data = analyze_with_pipeline(request.text, request.context.dict() if request.context else None, request.custom_codes)
        return result_data
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

from fastapi import UploadFile, File
from document_parser import process_document
import os
import tempfile

@app.post("/upload/document")
async def upload_document(file: UploadFile = File(...)):
    """
    Upload PDF or Word document, parse structure, and return processed data
    """
    try:
        # Validate file type
        file_ext = file.filename.split('.')[-1].lower()
        if file_ext not in ['pdf', 'docx', 'doc']:
            raise HTTPException(status_code=400, detail="Only PDF and Word documents are supported")
        
        # Save temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}") as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        # Process document
        result = process_document(tmp_path, file_ext)
        
        # Clean up
        os.unlink(tmp_path)
        
        return {
            "filename": file.filename,
            "processed_data": result,
            "message": "Document processed successfully"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/import/qdpx")
async def import_qdpx(file: UploadFile = File(...)):
    """
    Upload and parse a .qdpx (Atlas.ti) project file to extract codes.
    """
    try:
        # Validate file type
        if not file.filename.endswith('.qdpx'):
            raise HTTPException(status_code=400, detail="Only .qdpx files are supported")
        
        # Save temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=".qdpx") as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        # Extract codes
        codes = extract_codes_from_qdpx(tmp_path)
        
        # Clean up
        os.unlink(tmp_path)
        
        return {
            "filename": file.filename,
            "codes": codes,
            "count": len(codes),
            "message": "Codes extracted successfully"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

import json
from fastapi import Form

import io
import pypdf
from docx import Document
from typing import List

@app.post("/analyze/document")
async def analyze_document(
    files: List[UploadFile] = File(...),
    protocol: Optional[UploadFile] = File(None),
    context: Optional[str] = Form(None)
):
    try:
        combined_text = ""
        
        # Process protocol if provided
        protocol_text = ""
        if protocol:
            content = await protocol.read()
            if protocol.filename.endswith(".pdf"):
                pdf_reader = pypdf.PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    protocol_text += page.extract_text() + "\n"
            elif protocol.filename.endswith(".docx"):
                doc = Document(io.BytesIO(content))
                for para in doc.paragraphs:
                    protocol_text += para.text + "\n"
            else:
                protocol_text = content.decode("utf-8")

        # Process interview files
        for file in files:
            content = await file.read()
            file_text = ""
            if file.filename.endswith(".pdf"):
                pdf_reader = pypdf.PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    file_text += page.extract_text() + "\n"
            elif file.filename.endswith(".docx"):
                doc = Document(io.BytesIO(content))
                for para in doc.paragraphs:
                    file_text += para.text + "\n"
            else:
                file_text = content.decode("utf-8")
            
            combined_text += f"\n--- INTERVIEW: {file.filename} ---\n{file_text}\n"
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class SynthesisRequest(BaseModel):
    analyses: List[Dict[str, Any]]

@app.post("/analyze/synthesis")
def create_synthesis(request: SynthesisRequest):
    """
    Perform Cross-Case Synthesis on a list of completed analyses.
    """
    try:
        result = synthesize_structure(request.analyses)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
        if not combined_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from files")

        # Parse context JSON
        context_dict = {}
        if context:
            try:
                context_dict = json.loads(context)
                if protocol_text:
                    context_dict["interview_protocol"] = protocol_text
            except json.JSONDecodeError:
                pass

        analysis_result = analyze_with_pipeline(combined_text, context_dict)
        
        return {"filename": "multi-file-analysis", "analysis": analysis_result}
    except Exception as e:
        print(f"Error processing file: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))



class ComparativeRequest(BaseModel):
    texts: List[str]

@app.post("/analyze/comparative")
def create_comparative_analysis(request: ComparativeRequest):
    try:
        result_text = perform_comparative_analysis(request.texts)
        return {"result": result_text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/analyses")
def get_analyses(skip: int = 0, limit: int = 10, db: Session = Depends(get_db)):
    return db.query(Analysis).offset(skip).limit(limit).all()

import glob

@app.post("/demo/generate")
def generate_demo_results():
    """
    Returns aggregated results from the batch processed interviews.
    """
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        project_root = os.path.dirname(base_dir)
        
        # 1. Load Context
        context_path = os.path.join(project_root, "data", "demo", "context.json")
        context_data = {}
        if os.path.exists(context_path):
            with open(context_path, 'r') as f:
                context_data = json.load(f)

        # 2. Load Analysis Results
        results_dir = os.path.join(project_root, "analysis_results")
        analysis_files = glob.glob(os.path.join(results_dir, "*.json"))
        
        aggregated_codes = []
        dim_stats = {}
        processed_count = 0
        
        for file_path in analysis_files:
            try:
                with open(file_path, 'r') as f:
                    data = json.load(f)
                    
                    # Extract codes
                    # Structure might vary slightly, usually data['codes'] or data['phase1_codes']['codes']
                    codes = data.get('codes', [])
                    if not codes and 'phase1_codes' in data:
                        codes = data['phase1_codes'].get('codes', [])
                        
                    # Add participant ID to codes if missing
                    pid = os.path.basename(file_path).replace('.json', '')
                    for code in codes:
                        code['participant_id'] = pid
                        
                    aggregated_codes.extend(codes)
                    
                    # Aggregate stats
                    # data['dimensional_statistics']
                    stats = data.get('dimensional_statistics', {})
                    for dim, stat in stats.items():
                        if dim not in dim_stats:
                            dim_stats[dim] = {"total_codes": 0}
                        dim_stats[dim]["total_codes"] += stat.get("total_codes", 0)
                        
                    processed_count += 1
            except Exception as e:
                print(f"Error reading {file_path}: {e}")

        # Construct response
        analysis_result = {
            "participant_id": "BATCH_DEMO",
            "phenomenon_nucleus": f"Batch analysis of {processed_count} interviews. Detailed codes are aggregated below.",
            "codes": aggregated_codes,
            "dimensional_statistics": dim_stats,
            # Add other phases if we want valid demo data for them
            "markdown_table": f"| Metric | Value |\n|---|---|\n| Processed Interviews | {processed_count} |\n| Total Codes | {len(aggregated_codes)} |"
        }
        
        return {
            "context": context_data,
            "analysis": analysis_result
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
